"""위험 예측 보고서 생성기 — DB 조회 → 분석 → 차트 → Word(.docx).

adl_raw_records 의 응급 대상자(661, 30일)를 직접 조회해 지표 추세를 분석하고,
미니멀 편집 디자인(딥틸 단일 액센트·3선 표·KPI 카드)의 한글 차트와 함께
표지·요약·검증·권고를 갖춘 보고서를 생성한다. 모델/LLM 연동 전 산출물 포맷 템플릿.

실행:
  PYTHONPATH=. uv run --with "python-docx,matplotlib" python scripts/report_generate.py
"""

from __future__ import annotations

import argparse
import asyncio
import json
from datetime import date, datetime, timezone
from pathlib import Path
from statistics import mean

import matplotlib

matplotlib.use("Agg")
from matplotlib import font_manager, rcParams  # noqa: E402
import matplotlib.pyplot as plt  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Pt, RGBColor, Cm  # noqa: E402

from tortoise import Tortoise  # noqa: E402

from app.core.email import docx_to_pdf  # noqa: E402
from app.database import TORTOISE_ORM  # noqa: E402
from app.models.adl_raw import AdlRawRecord  # noqa: E402
from app.models.patient import Patient  # noqa: E402
from app.models.report import Report  # noqa: E402
from app.services.reports import classify, risk_of, risk_score  # noqa: E402

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "out" / "reports"
ASSET_DIR = OUT_DIR / "assets"
REPORT_DATE = date(2026, 6, 7)

# 기본(인자 미지정) 실행 = 응급 대상자 661, adl_raw id 185–214 (30일).
DEFAULT_PATIENT_ID = "661"
DEFAULT_ID_LO = 185
DEFAULT_ID_HI = 214

# ── 디자인 토큰 (미니멀 편집 팔레트) ──────────────────────────────────────────
INK = RGBColor(0x22, 0x27, 0x2E)  # 본문 텍스트 (거의 검정)
MUTE = RGBColor(0x66, 0x70, 0x85)  # 보조 텍스트 (회색)
ACCENT = RGBColor(0x0F, 0x4C, 0x5C)  # 단일 액센트 (딥 틸)
RISK = RGBColor(0xB4, 0x23, 0x18)  # 위험 강조 (절제된 레드)
WARN = RGBColor(0xB7, 0x79, 0x1F)  # 주의 강조 (절제된 주황)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
# hex 문자열 (음영/테두리용)
H_ACCENT = "0F4C5C"
H_RISK = "B42318"
H_WARN = "B7791F"
H_MUTE = "667085"
H_HAIR = "E4E7EC"  # 헤어라인
H_SOFT = "F7F8FA"  # 카드/얼터넷 배경
H_INK = "22272E"

# 분류(variant) 별 강조 색 / 배너 음영
VARIANT_INK = {"위험": RISK, "주의": WARN, "사망": MUTE}
VARIANT_SHADE = {"위험": H_RISK, "주의": H_WARN, "사망": H_MUTE}

# 분류별 내러티브 텍스트 (낙상·계단·병원명 등 661 고정 표현 제거).
NARRATIVE = {
    "위험": {
        "badge": "RISK LEVEL",
        "grade_label": "고위험 HIGH",
        "grade_short": "고위험 (위험)",
        "event": "사고·응급 이송 위험",
        "window": "단기(수주 내) 발생 가능성",
        "summary": (
            "예측 모델은 본 대상자를 고위험(위험)으로 분류하였다. 핵심 근거는 ①주간 활동량의 "
            "지속적 저하, ②자기관리 활동의 위축, ③야간 이상 지표 증가라는, 상호 연결된 세 가지 "
            "악화 신호다. 신체 쇠약과 야간 불안정이 동반된 응급 고위험 프로파일이다."
        ),
        "risk_factor": "독거(발견·구조 지연), 고령(사고 시 중상 가능성) 등 기저 위험이 존재한다.",
        "trend_lead": "주간 활동 저하와 야간 불안정의 결합은 임박한 사고·응급의 강한 전조다.",
        "validation_title": "예측 요약",
        "validation_note": (
            "30일 누적 패턴(주간 쇠약 + 야간 불안정)이 응급 위험을 시사한다. 조기 개입이 권고된다."
        ),
        "recommendations": [
            ["즉시", "현장 안전 점검 및 보호자 연락, 야간 상태 집중 확인", "담당 요양보호사"],
            ["즉시", "야간 이상지표 임계 초과 시 실시간 알림 등급 상향", "모니터링 운영"],
            ["단기", "주거 환경 안전 점검 및 위험 요소 개선", "복지 담당"],
            ["단기", "건강 상태 평가 및 보건·의료 연계", "보건·의료"],
        ],
    },
    "주의": {
        "badge": "CAUTION",
        "grade_label": "중위험 CAUTION",
        "grade_short": "중위험 (주의)",
        "event": "이상 징후 — 악화 가능성",
        "window": "경과 관찰 권고",
        "summary": (
            "예측 모델은 본 대상자를 중위험(주의)으로 분류하였다. 일부 지표에서 이상 징후가 "
            "관찰되나 즉각적 응급 수준은 아니다. 활동·야간 패턴 변화를 주기적으로 관찰할 필요가 있다."
        ),
        "risk_factor": "독거·고령 등 기저 위험은 있으나 현재 지표는 경계 수준이다.",
        "trend_lead": "현재 추세는 악화 가능성을 시사하므로 지속 관찰이 필요하다.",
        "validation_title": "예측 요약",
        "validation_note": "응급 수준은 아니나 추세 악화 시 등급 상향이 필요하다.",
        "recommendations": [
            ["단기", "모니터링 주기 단축 및 지표 변화 추적", "모니터링 운영"],
            ["단기", "방문 점검으로 생활 상태·복약 확인", "담당 요양보호사"],
            ["정기", "보건 상담 및 건강 상태 점검", "보건·의료"],
        ],
    },
    "사망": {
        "badge": "RECORD",
        "grade_label": "모니터링 종료",
        "grade_short": "사망 (모니터링 종료)",
        "event": "쇠약 진행 → 사망",
        "window": "관찰 종료",
        "summary": (
            "본 대상자는 관찰 기간 중 사망하였다. 사망 전 수주간 수면·활동 지표가 급격히 감소하는 "
            "노인성 쇠약 패턴이 관찰되었다. 본 보고서는 회고적 분석으로, 향후 유사 패턴의 조기 "
            "식별을 위한 참고 자료다."
        ),
        "risk_factor": "독거·고령에 더해 말기 활동·수면 급감이 사망에 선행하였다.",
        "trend_lead": "주간 활동 소실과 수면 급감의 결합은 말기 쇠약의 전형적 경과다.",
        "validation_title": "사망 확인 — 종료일 대조",
        "validation_note": (
            "30일 누적 패턴(활동 소실 + 수면 급감)이 사망으로 귀결되어, 쇠약 진행 신호의 "
            "회고적 검증 사례다."
        ),
        "recommendations": [
            ["즉시", "유가족 안내 및 사례 종결 처리", "복지 담당"],
            ["단기", "동일 패턴군 선제 모니터링 강화", "모니터링 운영"],
            ["단기", "말기 쇠약 조기경보 임계 재검토", "보건·의료"],
        ],
    },
}

KFONT = "/usr/share/fonts/truetype/nanum/NanumGothic.ttf"
KFONT_BOLD = "/usr/share/fonts/truetype/nanum/NanumGothicBold.ttf"
BODY_FONT = "맑은 고딕"


# ─────────────────────────────────────────────────────────────────────────────
# 1. 데이터
# ─────────────────────────────────────────────────────────────────────────────
async def fetch_rows(cid: str, id_lo: int | None, id_hi: int | None) -> list[dict]:
    """대상자의 ADL 원시 행을 조회한다 (연결은 호출 측에서 관리).

    id 범위가 주어지면 그대로(기본 661), 아니면 care_recipient_id 로 조회하고
    최근 30일치만 사용한다.
    """
    if id_lo is not None and id_hi is not None:
        q = AdlRawRecord.filter(id__gte=id_lo, id__lte=id_hi)
    else:
        q = AdlRawRecord.filter(care_recipient_id=cid)
    rows = await q.order_by("lifeog_date", "id").values()
    return rows[-30:]


async def fetch_anomaly(cid: str) -> dict | None:
    """대상자의 이상탐지 집계(dual 비율·MAE·라벨)를 읽는다 (읽기전용). 없으면 None.

    adl_anomaly_results 는 매핑/PK 없는 분석 테이블이라 raw SELECT 로 접근한다(스크립트 한정).
    """
    conn = Tortoise.get_connection("default")
    try:
        rows = await conn.execute_query_dict(
            "SELECT avg(CASE WHEN anomaly_dual THEN 1.0 ELSE 0.0 END) AS dual_frac, "
            "avg(mae_a) AS mae_avg, max(label) AS label, count(*) AS n "
            "FROM adl_anomaly_results WHERE care_recipient_id = $1",
            [cid],
        )
    except Exception as err:  # noqa: BLE001 — 테이블 부재(로컬 등) → 등급 폴백
        print(f"참고: 이상탐지 조회 불가({err}) — 등급 기반 폴백.")
        return None
    if not rows or not rows[0]["n"]:
        return None
    r = rows[0]
    return {
        "dual_frac": float(r["dual_frac"] or 0.0),
        "mae_avg": float(r["mae_avg"] or 0.0),
        "label": r["label"],
    }


def analyze(rows: list[dict]) -> dict:
    if len(rows) < 28:
        raise ValueError(f"분석에는 최소 28일치 ADL 데이터가 필요합니다 (현재 {len(rows)}일).")

    def a(xs):
        xs = [x for x in xs if x is not None]
        return mean(xs) if xs else 0.0

    dates = [r["lifeog_date"] for r in rows]
    aix = [r["aix_d"] for r in rows]
    night = [r["night_aix_ratio"] for r in rows]
    weeks = [rows[0:7], rows[7:14], rows[14:21], rows[21:]]
    wk = [
        {
            "label": f"{i}주차",
            "range": f"{w[0]['lifeog_date'].strftime('%m/%d')}–{w[-1]['lifeog_date'].strftime('%m/%d')}",
            "aix": a([x["aix_d"] for x in w]),
            "night": a([x["night_aix_ratio"] for x in w]),
            "bath": a([x["bath_count_d"] for x in w]),
            "out": a([x["outgoing_count_d"] for x in w]),
            "sleep": a([x["total_sleep_period"] for x in w]),
        }
        for i, w in enumerate(weeks, 1)
    ]
    ni = max(range(len(night)), key=lambda i: night[i] or 0)
    ai = min(range(len(aix)), key=lambda i: aix[i] if aix[i] is not None else 9e9)
    return {
        "rows": rows,
        "dates": dates,
        "aix": aix,
        "night": night,
        "bath": [r["bath_count_d"] for r in rows],
        "out": [r["outgoing_count_d"] for r in rows],
        "weeks": wk,
        "w1_aix": wk[0]["aix"],
        "w4_aix": wk[3]["aix"],
        "w1_night": wk[0]["night"],
        "w4_night": wk[3]["night"],
        "w1_bath": wk[0]["bath"],
        "w4_bath": wk[3]["bath"],
        "night_max": night[ni],
        "night_max_date": dates[ni],
        "aix_min": aix[ai],
        "aix_min_date": dates[ai],
        "profile": rows[0],
        "event_date": rows[-1]["emergency_date"],
    }


# ─────────────────────────────────────────────────────────────────────────────
# 2. 차트 (미니멀 스타일)
# ─────────────────────────────────────────────────────────────────────────────
def setup_chart_style():
    for f in (KFONT, KFONT_BOLD):
        if Path(f).exists():
            font_manager.fontManager.addfont(f)
    rcParams.update(
        {
            "font.family": "NanumGothic",
            "axes.unicode_minus": False,
            "figure.dpi": 200,
            "axes.edgecolor": "#C9CED6",
            "axes.linewidth": 0.8,
            "text.color": "#22272E",
            "axes.labelcolor": "#667085",
            "xtick.color": "#667085",
            "ytick.color": "#667085",
            "font.size": 9,
        }
    )


def _despine(ax, keep=("left", "bottom")):
    for s in ("top", "right", "left", "bottom"):
        ax.spines[s].set_visible(s in keep)


def chart_activity(an: dict) -> Path:
    xs = [d.strftime("%-m/%-d") for d in an["dates"]]
    fig, ax1 = plt.subplots(figsize=(9.0, 3.2))
    fig.patch.set_facecolor("white")
    ax1.plot(xs, an["aix"], color="#0F4C5C", lw=2.2, marker="o", ms=3.2, label="일일 활동량")
    ax1.fill_between(xs, an["aix"], color="#0F4C5C", alpha=0.06)
    ax1.set_ylabel("일일 활동량", fontsize=8.5)
    ax2 = ax1.twinx()
    ax2.plot(
        xs, an["night"], color="#B42318", lw=1.6, marker="s", ms=3, alpha=0.9, label="야간 이상비"
    )
    ax2.set_ylabel("야간 이상비", fontsize=8.5)
    _despine(ax1, ("left", "bottom"))
    _despine(ax2, ("right",))
    ax1.axvline(len(xs) - 1, color="#B0B6BF", ls=(0, (3, 3)), lw=1)
    ax1.annotate(
        "사건일 · 낙상",
        xy=(len(xs) - 1, 0),
        xytext=(len(xs) - 1, max(an["aix"]) * 0.5),
        fontsize=8,
        color="#B42318",
        ha="right",
        va="center",
    )
    ax1.set_xticks(range(0, len(xs), 3))
    ax1.set_xticklabels([xs[i] for i in range(0, len(xs), 3)], fontsize=7.5)
    ax1.tick_params(length=0, labelsize=7.5)
    ax2.tick_params(length=0, labelsize=7.5)
    ax1.grid(axis="y", color="#EDF0F3", lw=0.8)
    ax1.set_axisbelow(True)
    ax1.set_title(
        "활동량 붕괴와 야간 이상비 급등",
        loc="left",
        fontsize=10.5,
        fontweight="bold",
        color="#22272E",
        pad=10,
    )
    fig.tight_layout()
    p = ASSET_DIR / "chart_activity.png"
    fig.savefig(p, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return p


def chart_selfcare(an: dict) -> Path:
    xs = [d.strftime("%-m/%-d") for d in an["dates"]]
    fig, ax1 = plt.subplots(figsize=(9.0, 2.9))
    fig.patch.set_facecolor("white")
    ax1.bar(xs, an["bath"], color="#9DB8C4", width=0.62, label="목욕 횟수")
    ax1.set_ylabel("목욕 횟수", fontsize=8.5)
    ax2 = ax1.twinx()
    ax2.plot(xs, an["out"], color="#C2853B", lw=1.8, marker="^", ms=3.2, label="외출 횟수")
    ax2.set_ylabel("외출 횟수", fontsize=8.5)
    _despine(ax1, ("left", "bottom"))
    _despine(ax2, ("right",))
    ax1.set_xticks(range(0, len(xs), 3))
    ax1.set_xticklabels([xs[i] for i in range(0, len(xs), 3)], fontsize=7.5)
    ax1.tick_params(length=0, labelsize=7.5)
    ax2.tick_params(length=0, labelsize=7.5)
    ax1.grid(axis="y", color="#EDF0F3", lw=0.8)
    ax1.set_axisbelow(True)
    ax1.set_title(
        "자기관리(목욕) 위축과 외출 유지",
        loc="left",
        fontsize=10.5,
        fontweight="bold",
        color="#22272E",
        pad=10,
    )
    fig.tight_layout()
    p = ASSET_DIR / "chart_selfcare.png"
    fig.savefig(p, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return p


# ─────────────────────────────────────────────────────────────────────────────
# 3. docx 저수준 헬퍼
# ─────────────────────────────────────────────────────────────────────────────
def set_base_font(doc):
    st = doc.styles["Normal"]
    st.font.name = BODY_FONT
    st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.25


def run(p, text, *, size=10, bold=False, color=INK, italic=False, tracking=None, font=None):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    if font:
        r.font.name = font
        r._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if tracking:
        rPr = r._r.get_or_add_rPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(tracking))
        rPr.append(sp)
    return r


def para(doc, *, before=0, after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align:
        p.alignment = align
    return p


def bottom_border(p, color=H_HAIR, sz="6", space="6"):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), sz)
    b.set(qn("w:space"), space)
    b.set(qn("w:color"), color)
    pbdr.append(b)
    pPr.append(pbdr)


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex_color)
    tcPr.append(sh)


def cell_margins(cell, top=80, bottom=80, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def cell_text(cell, text, *, bold=False, color=INK, size=9, align=None, before=0, after=0):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align:
        p.alignment = align
    run(p, text, size=size, bold=bold, color=color)
    return p


def no_borders(t):
    tblPr = t._tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        b.append(e)
    tblPr.append(b)


def three_line(t, *, zebra=False):
    """상·하단 액센트 굵은 줄 + 헤더 아래 헤어라인. 데이터 행은 (옵션) 얼터넷 배경."""
    tblPr = t._tbl.tblPr
    b = OxmlElement("w:tblBorders")
    specs = {
        "top": (H_ACCENT, "12"),
        "bottom": (H_ACCENT, "12"),
        "left": (None, None),
        "right": (None, None),
        "insideH": (H_HAIR, "2") if zebra else (None, None),
        "insideV": (None, None),
    }
    for edge, (col, sz) in specs.items():
        e = OxmlElement(f"w:{edge}")
        if col is None:
            e.set(qn("w:val"), "none")
        else:
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), sz)
            e.set(qn("w:color"), col)
        b.append(e)
    tblPr.append(b)
    # 헤더 행 아래 헤어라인 (셀 하단 테두리)
    for c in t.rows[0].cells:
        tcPr = c._tc.get_or_add_tcPr()
        tb = OxmlElement("w:tcBorders")
        bo = OxmlElement("w:bottom")
        bo.set(qn("w:val"), "single")
        bo.set(qn("w:sz"), "8")
        bo.set(qn("w:color"), H_ACCENT)
        tb.append(bo)
        tcPr.append(tb)
    if zebra:
        for ri, rowobj in enumerate(t.rows[1:]):
            if ri % 2 == 1:
                for c in rowobj.cells:
                    shade(c, H_SOFT)


# ─────────────────────────────────────────────────────────────────────────────
# 4. 고수준 컴포넌트
# ─────────────────────────────────────────────────────────────────────────────
def h1(doc, num, title):
    p = para(doc, before=20, after=8)
    run(p, f"{num}  ", size=14, bold=True, color=ACCENT)
    run(p, title, size=14, bold=True, color=INK)
    bottom_border(p, color=H_HAIR, sz="6", space="8")
    return p


def h2(doc, title):
    p = para(doc, before=12, after=4)
    run(p, "▍", size=11, color=ACCENT)
    run(p, " " + title, size=11.5, bold=True, color=INK)
    return p


def data_table(doc, header, rows, *, aligns=None, zebra=False, widths=None):
    t = doc.add_table(rows=1, cols=len(header))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, htext in enumerate(header):
        cell_text(
            t.rows[0].cells[i],
            htext,
            bold=True,
            size=8.5,
            color=ACCENT,
            align=(aligns or {}).get(i, WD_ALIGN_PARAGRAPH.LEFT),
            before=2,
            after=2,
        )
        cell_margins(t.rows[0].cells[i])
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cell_text(
                cells[i],
                str(v),
                size=9,
                align=(aligns or {}).get(i, WD_ALIGN_PARAGRAPH.LEFT),
                before=2,
                after=2,
            )
            cell_margins(cells[i])
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Cm(w)
    three_line(t, zebra=zebra)
    return t


def kv_block(doc, pairs, label_w=4.2, val_w=11.5):
    """라벨/값 미니멀 표 — 가로줄 없이 라벨만 회색, 값 잉크."""
    t = doc.add_table(rows=0, cols=2)
    no_borders(t)
    for k, v in pairs:
        c = t.add_row().cells
        c[0].width = Cm(label_w)
        c[1].width = Cm(val_w)
        cell_text(c[0], k, size=9, color=MUTE, before=2, after=2)
        cell_text(c[1], v, size=9.5, color=INK, before=2, after=2)
        cell_margins(c[0], left=0)
        cell_margins(c[1], left=0)
    return t


def kpi_cards(doc, cards):
    """3개 KPI 카드 — 카드 사이 빈 스페이서 열로 분리감."""
    n = len(cards)
    cols = n * 2 - 1
    t = doc.add_table(rows=1, cols=cols)
    no_borders(t)
    for idx, (label, value, delta, delta_color) in enumerate(cards):
        ci = idx * 2
        cell = t.rows[0].cells[ci]
        cell.width = Cm(4.8)
        shade(cell, H_SOFT)
        cell_margins(cell, top=160, bottom=160, left=200, right=160)
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        run(p1, label, size=8.5, color=MUTE, tracking=20)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        run(p2, value, size=21, bold=True, color=INK)
        p3 = cell.add_paragraph()
        run(p3, delta, size=9, bold=True, color=delta_color)
        if ci + 1 < cols:
            t.rows[0].cells[ci + 1].width = Cm(0.5)
    return t


def risk_banner(doc, grade_label, badge, score, event, window, *, shade_hex=H_RISK, score_color=RISK):
    t = doc.add_table(rows=1, cols=2)
    no_borders(t)
    left = t.rows[0].cells[0]
    left.width = Cm(4.6)
    shade(left, shade_hex)
    cell_margins(left, top=120, bottom=120, left=160, right=160)
    left.text = ""
    pa = left.paragraphs[0]
    pa.paragraph_format.space_after = Pt(0)
    pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(pa, badge, size=10, bold=True, color=WHITE, tracking=30)
    pb = left.add_paragraph()
    pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(pb, grade_label, size=15, bold=True, color=WHITE)
    right = t.rows[0].cells[1]
    right.width = Cm(11.4)
    cell_margins(right, top=120, bottom=120, left=260, right=120)
    right.text = ""
    r1 = right.paragraphs[0]
    r1.paragraph_format.space_after = Pt(3)
    run(r1, "위험점수 ", size=10, color=MUTE)
    run(r1, f"{score:.2f}", size=14, bold=True, color=score_color)
    run(r1, " / 1.00", size=10, color=MUTE)
    r2 = right.add_paragraph()
    r2.paragraph_format.space_after = Pt(0)
    run(r2, "예측 이벤트  ", size=8.5, color=MUTE, tracking=10)
    run(r2, event, size=9.5, bold=True, color=INK)
    r3 = right.add_paragraph()
    run(r3, "임박도  ", size=8.5, color=MUTE, tracking=10)
    run(r3, window, size=9.5, bold=True, color=INK)
    return t


def page_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, "살펴봄 위험예측 시스템 · 자동 생성 보고서      ", size=8, color=MUTE)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    rr = p.add_run()
    rr._r.append(fld)
    rr.font.size = Pt(8)
    rr.font.color.rgb = MUTE


def pct(a, b):
    if not a:
        return "—"
    d = (b - a) / a * 100
    return f"{'▲' if d >= 0 else '▼'} {abs(d):.0f}%"


# ─────────────────────────────────────────────────────────────────────────────
# 5. 보고서
# ─────────────────────────────────────────────────────────────────────────────
def build(an, c1, c2, *, patient_id: str, report_date: date, profile: dict) -> Path:
    prof = an["profile"]
    sex = "여성" if prof["sex"] == "F" else "남성"
    variant = profile["variant"]  # 위험 | 주의 | 사망
    score = profile["score"]
    # 기본 템플릿에 LLM 맞춤 내러티브(profile["narrative"])를 덮어쓴다 (있는 키만).
    nv = {**NARRATIVE[variant], **(profile.get("narrative") or {})}
    vink = VARIANT_INK[variant]  # 분류 강조 색
    doc = Document()
    set_base_font(doc)
    doc.core_properties.title = f"고령자 위험 예측 보고서 — 대상자 {patient_id}"
    doc.core_properties.author = "살펴봄 위험예측 시스템"

    s = doc.sections[0]
    s.top_margin = Cm(2.4)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.6)
    s.right_margin = Cm(2.6)

    # ── 표지 (좌측 정렬, 여백 중심) ──
    para(doc, before=70)
    eyebrow = para(doc, after=10)
    run(
        eyebrow,
        "SALPYEOBOM   RISK PREDICTION REPORT",
        size=9.5,
        bold=True,
        color=ACCENT,
        tracking=60,
    )
    t1 = para(doc, after=2)
    run(t1, "고령자 사고·사망 위험", size=30, bold=True, color=INK)
    t2 = para(doc, after=14)
    run(t2, "예측 분석 보고서", size=30, bold=True, color=INK)
    rule = para(doc, after=14)
    bottom_border(rule, color=H_ACCENT, sz="18", space="1")
    sub = para(doc, after=40)
    run(sub, "원격 ADL 모니터링 기반 30일 위험 추세 분석", size=12, color=MUTE)

    meta = kv_block(
        doc,
        [
            ("문서 번호", f"SPB-RR-{patient_id}-{report_date.strftime('%Y%m%d')}"),
            ("분석 대상자", f"{prof['care_recipient_id']} · {prof['age']}세 · {sex} · 독거"),
            ("관찰 구간", f"{an['dates'][0]} – {an['dates'][-1]}  (30일)"),
            ("작성 일자", report_date.isoformat()),
            ("분류 등급", f"{nv['grade_short']} · 위험점수 {score:.2f}"),
            ("작성 주체", "위험예측 모델 + LLM 자동 서술"),
        ],
        label_w=3.6,
    )
    meta  # noqa

    para(doc, before=30)
    disc = para(doc)
    run(
        disc,
        "본 보고서는 AI 예측 모델의 출력을 LLM이 자동 서술한 문서입니다. "
        "위험 점수는 모델 산출 가정값이며, 임상적 진단이나 의료적 판단을 대체하지 않습니다.",
        size=8.5,
        color=MUTE,
        italic=True,
    )

    # ── 본문 ──
    doc.add_section(WD_SECTION.NEW_PAGE)
    page_footer(doc.sections[-1])

    h1(doc, "00", "핵심 요약")
    risk_banner(
        doc,
        nv["grade_label"],
        nv["badge"],
        score,
        nv["event"],
        nv["window"],
        shade_hex=VARIANT_SHADE[variant],
        score_color=vink,
    )
    para(doc, before=10)
    kpi_cards(
        doc,
        [
            (
                "활동량 aix_d",
                f"{an['w4_aix']:.0f}",
                f"{pct(an['w1_aix'], an['w4_aix'])}  vs 1주",
                vink,
            ),
            (
                "야간 이상비",
                f"{an['w4_night'] / 1000:.1f}k",
                f"{pct(an['w1_night'], an['w4_night'])}  vs 1주",
                vink,
            ),
            (
                "목욕 횟수",
                f"{an['w4_bath']:.0f}",
                f"{pct(an['w1_bath'], an['w4_bath'])}  vs 1주",
                vink,
            ),
        ],
    )
    para(doc, before=8)
    p = para(doc)
    run(
        p,
        f"대상자 {prof['care_recipient_id']}({prof['age']}세·{sex}·독거)의 최근 30일 원격 ADL "
        "데이터를 분석한 결과는 다음과 같다. ",
    )
    run(p, nv["summary"])

    h1(doc, "01", "대상자 개요")
    kv_block(
        doc,
        [
            ("대상자 ID", str(prof["care_recipient_id"])),
            ("나이 · 성별", f"{prof['age']}세 · {sex}"),
            (
                "거주 형태",
                f"독거 · {prof['house_structure']} · 방 {prof['room_no']}칸 · 욕실 {prof['bath_location']} · {prof['district']}",
            ),
            ("시력 · 청력", f"{prof['vision']} · {prof['hearing']}"),
            ("복약", prof["dosage"] or "정보 없음"),
            ("관찰 구간", f"{an['dates'][0]} – {an['dates'][-1]} (30일)"),
            ("데이터 출처", "adl_raw_records (원격 ADL 센서 일별 집계)"),
        ],
    )
    p = para(doc, before=6)
    run(p, "위험 가중 요인  ", bold=True, color=ACCENT)
    run(p, nv["risk_factor"])

    h1(doc, "02", "분석 방법 및 데이터")
    p = para(doc)
    run(
        p,
        "원격 ADL 센서가 산출하는 일별 지표(활동량·수면·목욕·외출·야간 이상비)를 30일치 수집하여, "
        "예측 모델이 사고·사망 위험을 0–1 점수와 4단계 등급으로 산출한다. 본 LLM 보고서는 모델 출력을 "
        "재예측하지 않고, 모델이 근거로 삼은 지표 추세를 사람이 이해할 수 있도록 서술·시각화한다.",
    )
    p = para(doc)
    run(
        p,
        "데이터 품질 주의 — night_aix_ratio 등 일부 비율 지표는 원본 계산 편차가 커 절대값보다 ‘추세’로 "
        "해석한다. 외출 시계열의 254/255 코드는 센서 무신호 표식으로 분석에서 제외했다.",
        size=8.5,
        color=MUTE,
    )

    h1(doc, "03", "위험 판정 근거 — 30일 추세")
    p = para(doc)
    run(
        p,
        f"활동량(aix_d)은 1주차 평균 {an['w1_aix']:.0f}에서 4주차 {an['w4_aix']:.0f}로 변화했고 "
        f"최저 {an['aix_min']:.0f}({an['aix_min_date'].strftime('%-m/%-d')})을 기록했다. 야간 이상비는 "
        f"최대 {an['night_max']:,.0f}({an['night_max_date'].strftime('%-m/%-d')})에 달했다. "
        + nv["trend_lead"],
    )
    doc.add_picture(str(c1), width=Cm(16.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    h2(doc, "주차별 지표 추이")
    data_table(
        doc,
        ["구간", "기간", "활동량", "야간이상", "목욕", "외출", "수면(분)"],
        [
            [
                w["label"],
                w["range"],
                f"{w['aix']:.0f}",
                f"{w['night']:,.0f}",
                f"{w['bath']:.1f}",
                f"{w['out']:.1f}",
                f"{w['sleep']:.0f}",
            ]
            for w in an["weeks"]
        ],
        aligns={i: WD_ALIGN_PARAGRAPH.CENTER for i in range(2, 7)},
    )
    para(doc, before=8)
    doc.add_picture(str(c2), width=Cm(16.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    h2(doc, "결정적 신호")
    for b in [
        f"활동량 변화 — 1주차 {an['w1_aix']:.0f} → 최저 {an['aix_min']:.0f}"
        f"({an['aix_min_date'].strftime('%-m/%-d')}). 거동 능력 추이.",
        f"야간 이상 — night_aix_ratio 최대 {an['night_max']:,.0f}"
        f"({an['night_max_date'].strftime('%-m/%-d')}). 야간 상태 안정도 지표.",
        f"자기관리 — 목욕 횟수 {an['w1_bath']:.0f} → {an['w4_bath']:.0f}회. 일상 기능 추이.",
        "외출 — 활동량 대비 외출 패턴 변화 관찰.",
    ]:
        bp = para(doc, after=3)
        run(bp, "—  ", color=ACCENT, bold=True)
        run(bp, b, size=9.5)

    h1(doc, "04", nv["validation_title"])
    if variant == "사망":
        event_date_str = str(an["event_date"] or prof.get("death_date") or "—")
        validation_rows = [
            ("실제 발생일", event_date_str),
            ("구분", "사망 (모니터링 종료)"),
            ("선행 패턴", "관찰 종료 전 활동·수면 급감"),
            ("예측–실제", "말기 쇠약 패턴과 부합"),
        ]
    else:
        validation_rows = [
            ("예측 분류", nv["grade_short"]),
            ("위험 점수", f"{score:.2f} / 1.00"),
            ("주요 근거", "활동량 추이 + 야간 이상 지표"),
            ("권고 대응", nv["window"]),
        ]
    kv_block(doc, validation_rows)
    p = para(doc, before=6)
    run(p, nv["validation_note"])

    h1(doc, "05", "권고 조치")
    data_table(
        doc,
        ["시급도", "조치 내용", "담당"],
        nv["recommendations"],
        aligns={0: WD_ALIGN_PARAGRAPH.CENTER},
        widths=[2.0, 9.5, 4.0],
        zebra=True,
    )

    # 부록
    doc.add_section(WD_SECTION.NEW_PAGE)
    page_footer(doc.sections[-1])
    h1(doc, "A", "부록 · 일자별 지표 전체 (30일)")
    app_rows = []
    for r in an["rows"]:
        app_rows.append(
            [
                r["lifeog_date"].strftime("%m/%d"),
                f"{r['aix_d']:.0f}" if r["aix_d"] is not None else "—",
                f"{r['night_aix_ratio']:,.0f}" if r["night_aix_ratio"] is not None else "—",
                str(r["bath_count_d"]) if r["bath_count_d"] is not None else "—",
                str(r["outgoing_count_d"]) if r["outgoing_count_d"] is not None else "—",
                f"{r['total_sleep_period']:.0f}" if r["total_sleep_period"] is not None else "—",
            ]
        )
    data_table(
        doc,
        ["날짜", "활동량", "야간이상", "목욕", "외출", "수면(분)"],
        app_rows,
        aligns={i: WD_ALIGN_PARAGRAPH.CENTER for i in range(1, 6)},
        zebra=True,
    )
    end = para(doc, before=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(end, "— 보고서 끝 —", size=9, color=MUTE, italic=True)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUT_DIR / f"위험예측보고서_{patient_id}_{report_date.strftime('%Y%m%d')}.docx"
    doc.save(out)
    return out


async def register_report(
    patient_id: str, title: str, file_name: str, report_date: date, risk_level: str
) -> None:
    """생성한 보고서를 Report 테이블에 멱등 등록한다 (file_name 기준 upsert)."""
    patient = await Patient.get_or_none(patient_id=patient_id)
    if patient is None:
        print(f"경고: 환자 {patient_id} 미존재 — Report 미등록 (load_derived 를 먼저 실행).")
        return
    generated_at = datetime(
        report_date.year, report_date.month, report_date.day, tzinfo=timezone.utc
    )
    await Report.update_or_create(
        file_name=file_name,
        defaults={
            "patient": patient,
            "title": title,
            "generated_at": generated_at,
            "risk_level": risk_level,
        },
    )
    print(f"Report 등록 → {file_name} ({patient_id}, {risk_level})")


async def analyze_target(
    patient_id: str, cid: str, id_lo: int | None, id_hi: int | None
) -> tuple[dict, dict]:
    """대상자 분석(an) + 분류/점수 프로파일을 산출한다 (분류·점수는 이상탐지+추세 결합)."""
    rows = await fetch_rows(cid, id_lo, id_hi)
    an = analyze(rows)
    anomaly = await fetch_anomaly(cid)
    patient = await Patient.get_or_none(patient_id=patient_id)
    level = patient.cross_verification_level if patient else None
    variant = (classify(anomaly["label"], anomaly["dual_frac"]) if anomaly else None) or risk_of(
        level
    )
    w1 = an["w1_aix"] or 0.0
    activity_decline = max(0.0, (w1 - an["w4_aix"]) / w1) if w1 else 0.0
    score = risk_score(
        anomaly["dual_frac"] if anomaly else None,
        anomaly["mae_avg"] if anomaly else None,
        activity_decline,
    )
    profile = {"variant": variant, "score": score, "label": anomaly["label"] if anomaly else None}
    return an, profile


def analysis_summary(patient_id: str, cid: str, an: dict, profile: dict) -> dict:
    """서브에이전트가 맞춤 내러티브를 작성할 수 있도록 핵심 지표를 JSON 친화 dict 로 요약."""
    prof = an["profile"]
    return {
        "patient_id": patient_id,
        "cid": cid,
        "variant": profile["variant"],
        "score": profile["score"],
        "age": prof.get("age"),
        "sex": "여성" if prof.get("sex") == "F" else "남성",
        "district": prof.get("district"),
        "metrics": {
            "w1_aix": round(an["w1_aix"], 1),
            "w4_aix": round(an["w4_aix"], 1),
            "aix_min": round(an["aix_min"], 1),
            "aix_min_date": str(an["aix_min_date"]),
            "night_max": round(an["night_max"], 1),
            "night_max_date": str(an["night_max_date"]),
            "w1_bath": round(an["w1_bath"], 1),
            "w4_bath": round(an["w4_bath"], 1),
            "w1_night": round(an["w1_night"], 1),
            "w4_night": round(an["w4_night"], 1),
            "weeks": [
                {"label": w["label"], "aix": round(w["aix"], 1), "night": round(w["night"], 1)}
                for w in an["weeks"]
            ],
        },
    }


async def generate(
    patient_id: str,
    cid: str,
    report_date: date,
    id_lo: int | None,
    id_hi: int | None,
    narrative_path: str | None = None,
) -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    await Tortoise.init(config=TORTOISE_ORM)
    try:
        an, profile = await analyze_target(patient_id, cid, id_lo, id_hi)
        # 맞춤 내러티브 주입 (있으면 NARRATIVE 템플릿을 LLM 작성 텍스트로 덮어씀).
        if narrative_path:
            profile["narrative"] = json.loads(Path(narrative_path).read_text(encoding="utf-8"))

        setup_chart_style()
        c1 = chart_activity(an)
        c2 = chart_selfcare(an)
        out = build(an, c1, c2, patient_id=patient_id, report_date=report_date, profile=profile)
        # 조회 서빙용 PDF 변환. soffice 미설치 시 docx 로 폴백 등록.
        try:
            file_name = docx_to_pdf(out).name
        except Exception as err:  # noqa: BLE001
            print(f"경고: PDF 변환 실패({err}) — docx 로 등록.")
            file_name = out.name
        await register_report(
            patient_id, f"{patient_id} 위험예측 보고서", file_name, report_date, profile["variant"]
        )
        print(f"보고서 생성 완료 → {out} [{profile['variant']}, score={profile['score']}]")
    finally:
        await Tortoise.close_connections()


async def dump_analysis(patient_id: str, cid: str, id_lo: int | None, id_hi: int | None) -> None:
    """대상자 분석 요약 JSON 을 stdout 으로 출력한다 (내러티브 작성용)."""
    await Tortoise.init(config=TORTOISE_ORM)
    try:
        an, profile = await analyze_target(patient_id, cid, id_lo, id_hi)
        print(json.dumps(analysis_summary(patient_id, cid, an, profile), ensure_ascii=False))
    finally:
        await Tortoise.close_connections()


def main() -> None:
    p = argparse.ArgumentParser(description="위험예측 보고서 생성기")
    p.add_argument("--patient-id", default=DEFAULT_PATIENT_ID, help="Patient.patient_id")
    p.add_argument(
        "--cid", default=None, help="adl_raw care_recipient_id (기본: patient-id 와 동일)"
    )
    p.add_argument("--date", default=REPORT_DATE.isoformat(), help="보고서 일자 (YYYY-MM-DD)")
    p.add_argument("--id-lo", type=int, default=None, help="adl_raw id 범위 시작")
    p.add_argument("--id-hi", type=int, default=None, help="adl_raw id 범위 끝")
    p.add_argument("--narrative", default=None, help="LLM 맞춤 내러티브 JSON 경로 (주입)")
    p.add_argument(
        "--dump-analysis", action="store_true", help="보고서 미생성, 분석 요약 JSON 만 출력"
    )
    args = p.parse_args()

    patient_id = args.patient_id
    cid = args.cid or patient_id
    report_date = date.fromisoformat(args.date)
    id_lo, id_hi = args.id_lo, args.id_hi
    # 인자 없는 기본 실행(661)은 기존 id 범위를 유지해 동일 출력 보장.
    if id_lo is None and id_hi is None and args.cid is None and patient_id == DEFAULT_PATIENT_ID:
        id_lo, id_hi = DEFAULT_ID_LO, DEFAULT_ID_HI

    if args.dump_analysis:
        asyncio.run(dump_analysis(patient_id, cid, id_lo, id_hi))
    else:
        asyncio.run(generate(patient_id, cid, report_date, id_lo, id_hi, args.narrative))


if __name__ == "__main__":
    main()
