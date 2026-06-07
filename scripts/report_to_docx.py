"""위험 예측 보고서(.md) → Word(.docx) 변환.

out/reports/risk_report_661.md 를 근거로 한국어 Word 보고서를 생성한다.
표·강조·색상 등 기본 서식을 적용한다. LLM/모델 연동 전, 보고서 산출물 포맷 확인용.

실행:
  uv run --with python-docx python scripts/report_to_docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "out" / "reports" / "risk_report_661.docx"

RED = RGBColor(0xC0, 0x00, 0x00)
GRAY = RGBColor(0x66, 0x66, 0x66)
BLUE = RGBColor(0x1F, 0x4E, 0x79)


def set_korean_font(doc: Document, name: str = "맑은 고딕") -> None:
    """기본 스타일에 한글 폰트를 지정 (Word 가 한글에 맞는 글꼴을 쓰도록)."""
    from docx.oxml.ns import qn

    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for k, v in rows:
        cells = t.add_row().cells
        cells[0].text = k
        cells[1].text = v
        cells[0].paragraphs[0].runs[0].bold = True


def add_grid_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v


def heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BLUE


def build() -> Path:
    doc = Document()
    set_korean_font(doc)

    # 제목
    title = doc.add_heading("고령자 위험 예측 보고서", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("응급·사고 조기경보 — 대상자 661 · 30일 관찰")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    sub.runs[0].font.color.rgb = GRAY

    note = doc.add_paragraph()
    r = note.add_run(
        "⚠️ 본 보고서는 AI 모델이 산출한 예측 결과를 LLM이 서술형으로 정리한 자동 생성 "
        "문서입니다. 위험 점수는 모델 산출 가정값(자리표시자)이며 임상 판단을 대체하지 않습니다."
    )
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY

    # 1. 대상자 개요
    heading(doc, "1. 대상자 개요")
    add_kv_table(
        doc,
        [
            ("대상자 ID", "661"),
            ("나이 / 성별", "78세 / 여성"),
            ("거주 형태", "독거 (도시·단독주택, 방 2칸, 옥내 욕실)"),
            ("시력 / 청력", "보통 / 보통"),
            ("복약", "없음"),
            ("관찰 구간", "2021-12-05 ~ 2022-01-03 (30일)"),
        ],
    )

    # 2. 종합 판정
    heading(doc, "2. 종합 판정")
    add_kv_table(
        doc,
        [
            ("예측 위험 등급", "고위험 (HIGH) · 4단계"),
            ("예측 위험 점수", "0.87 / 1.00 (모델 산출 가정값)"),
            ("예측 이벤트 유형", "낙상·외상에 의한 응급 이송"),
            ("임박도", "1주 이내 응급 발생 가능성 높음"),
        ],
    )
    p = doc.add_paragraph()
    p.add_run("요지 — ").bold = True
    p.add_run(
        "관찰 후반부로 갈수록 주간 활동량은 급감하고 야간 이상 지표(night_aix)는 폭증하는, "
        "전형적인 '쇠약 + 야간 불안정' 패턴이 누적되었다. 모델은 이를 임박한 낙상·응급 신호로 판정했다."
    )

    # 3. 핵심 근거
    heading(doc, "3. 핵심 근거 — 30일 추세")
    heading(doc, "3.1 지표별 초기 vs 말기 비교", level=2)
    add_grid_table(
        doc,
        ["지표", "초기 7일", "말기 7일", "변화", "해석"],
        [
            ["일일 활동(aix_d)", "166.9", "44.0", "▼74%", "활동성 급감 → 거동 저하"],
            [
                "야간 이상비(night_aix_ratio)",
                "3,893",
                "7,339",
                "▲89%",
                "야간 불안정·수면 중 이상 증가",
            ],
            ["목욕 횟수(bath_count_d)", "23.3", "7.4", "▼68%", "자기관리·활동 위축"],
            ["외출 횟수(outgoing_count_d)", "8.9", "7.7", "▼13%", "외출 유지(외부 사고 노출 지속)"],
            ["수면 시간(total_sleep_period)", "15분", "241분", "▲", "주간 무기력·수면 패턴 붕괴"],
        ],
    )
    heading(doc, "3.2 결정적 신호", level=2)
    for b in [
        "활동량 붕괴: aix_d 가 252(12/05) → 최저 4(12/21) 까지 떨어졌다 회복하지 못함.",
        "야간 이상치 폭발: night_aix_ratio 가 말기에 12/28: 13,113 → 12/29: 20,417(최대) 로 급등. "
        "사고 직전 야간 상태가 극도로 불안정했음을 시사.",
        "외출은 유지: 활동 전반은 줄었으나 외출 빈도는 유지 → 약화된 신체로 외부 보행을 지속하는 낙상 고위험 조합.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    # 4. 예측 vs 실제
    heading(doc, "4. 예측 결과 vs 실제 발생 (검증)")
    add_kv_table(
        doc,
        [
            ("실제 발생일", "2022-01-03"),
            ("발생 장소", "실외 (계단)"),
            ("경위", "계단에서 넘어져 119 호출 → 세일병원 이송·입원"),
            ("모델 예측 일치 여부", "일치 — 고위험 판정 구간 내 낙상·응급 이송 발생"),
        ],
    )
    p = doc.add_paragraph()
    p.add_run(
        "30일 누적 패턴(주간 쇠약 + 야간 불안정 + 외출 유지)이 실제 실외 계단 낙상으로 귀결되어, "
        "모델의 고위험 예측이 실제 이벤트와 부합했다."
    )

    # 5. 권고 조치
    heading(doc, "5. 권고 조치")
    for b in [
        "즉시 — 담당 요양보호사 현장 점검 및 보호자 연락 (야간 상태 집중 확인)",
        "환경 — 계단·욕실 등 실외/실내 동선 미끄럼 방지·안전손잡이 점검",
        "모니터링 — 야간 이상지표 임계 초과 시 실시간 알림 등급 상향",
        "의료 — 보행·근력 평가 및 낙상 예방 중재 연계",
    ]:
        doc.add_paragraph(b, style="List Number")

    # 푸터
    foot = doc.add_paragraph()
    fr = foot.add_run(
        "생성: 살펴봄(salpyeobom) 위험예측 파이프라인 · 본 문서는 자동 생성 초안이며 "
        "임상 판단을 대체하지 않습니다."
    )
    fr.italic = True
    fr.font.size = Pt(8)
    fr.font.color.rgb = GRAY

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Word 보고서 생성 완료 → {path}")
